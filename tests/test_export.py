import json
import zipfile
from io import BytesIO

import pytest
from docx import Document

from services.export.exporters import FORMATS, export_run

pytestmark = pytest.mark.unit

RUN = {
    "id": "abcdef1234567890",
    "direction": "zh-en",
    "status": "COMPLETED",
    "summary": "Test export",
    "confidentiality": "PUBLIC",
    "pipeline_version": "0.1.0",
    "version_pins": {},
}
SEGMENTS = [
    {
        "idx": 0,
        "source": "推动高质量发展。",
        "translation": "Promote high-quality development.",
        "versions": {"ai_draft": "Promote high-quality development."},
    },
    {
        "idx": 1,
        "source": "2023年增长5.2%。",
        "translation": "In 2023 it grew by 5.2%.",
        "versions": {"ai_draft": "In 2023 it grew by 5.2%.", "final": "In 2023 it grew by 5.2%."},
    },
]


class TestTextFormats:
    def test_txt(self):
        out = export_run(RUN, SEGMENTS, "txt")
        text = out.content.decode()
        assert "Promote high-quality development." in text
        assert out.filename.endswith(".txt")

    def test_md_is_clean_reader_facing_document(self):
        out = export_run(RUN, SEGMENTS, "md")
        text = out.content.decode()
        assert text.startswith("# Test export")
        assert "direction:" not in text and "pipeline:" not in text

    def test_json_roundtrip(self):
        out = export_run(RUN, SEGMENTS, "json")
        payload = json.loads(out.content)
        assert payload["segments"][1]["versions"]["final"] == "In 2023 it grew by 5.2%."


class TestInterchangeFormats:
    def test_xliff(self):
        out = export_run(RUN, SEGMENTS, "xliff")
        text = out.content.decode()
        assert '<xliff version="1.2"' in text
        assert '<source xml:lang="zh-CN">推动高质量发展。</source>' in text
        assert text.count("<trans-unit") == 2

    def test_tmx(self):
        out = export_run(RUN, SEGMENTS, "tmx")
        text = out.content.decode()
        assert '<tuv xml:lang="zh-CN">' in text
        assert "5.2%" in text

    def test_escaping(self):
        segs = [
            {
                "idx": 0,
                "source": '含<标签>与"引号"',
                "translation": 'with <tag> & "quote"',
                "versions": {},
            }
        ]
        text = export_run(RUN, segs, "xliff").content.decode()
        assert "&lt;标签&gt;" in text and "&quot;" in text

    def test_interchange_languages_follow_the_run_pair(self):
        run = {
            **RUN,
            "direction": "fr-de",
            "source_language": "fr",
            "target_language": "de",
        }
        xliff = export_run(run, SEGMENTS, "xliff").content.decode()
        tmx = export_run(run, SEGMENTS, "tmx").content.decode()
        assert 'source-language="fr-FR"' in xliff
        assert 'target-language="de-DE"' in xliff
        assert 'xml:lang="fr-FR"' in tmx
        assert 'xml:lang="de-DE"' in tmx


class TestDocx:
    def test_docx_structure(self):
        out = export_run(RUN, SEGMENTS, "docx")
        assert out.content[:2] == b"PK"  # zip magic
        with zipfile.ZipFile(BytesIO(out.content)) as zf:
            names = zf.namelist()
            assert "word/document.xml" in names
            document = zf.read("word/document.xml").decode("utf-8")
        assert "Promote high-quality development." in document
        assert "direction:" not in document and "pipeline:" not in document

    def test_bilingual_table(self):
        out = export_run(RUN, SEGMENTS, "docx_bilingual")
        with zipfile.ZipFile(BytesIO(out.content)) as zf:
            document = zf.read("word/document.xml").decode("utf-8")
        assert "推动高质量发展。" in document and "Promote high-quality development." in document
        assert "<w:tbl>" in document  # basic table preserved

    def test_arabic_docx_uses_dynamic_header_and_rtl_paragraphs(self):
        run = {
            **RUN,
            "direction": "en-ar",
            "source_language": "en",
            "target_language": "ar",
        }
        segments = [
            {
                "idx": 0,
                "source": "International cooperation supports development.",
                "translation": "يدعم التعاون الدولي التنمية.",
                "versions": {},
            }
        ]
        out = export_run(run, segments, "docx_bilingual")
        with zipfile.ZipFile(BytesIO(out.content)) as zf:
            document = zf.read("word/document.xml").decode("utf-8")
        assert "阿拉伯语译文" in document
        assert "<w:bidi" in document

    def test_docx_restores_reader_facing_title_and_layout(self):
        structured = [
            {
                "idx": 0,
                "source": "政府工作报告",
                "translation": "Government Work Report",
                "versions": {},
            },
            {
                "idx": 1,
                "source": "一、总体要求",
                "translation": "I. General Requirements",
                "versions": {},
            },
            {
                "idx": 2,
                "source": "扎实推进各项工作。",
                "translation": "We will advance all tasks.",
                "versions": {},
            },
        ]
        out = export_run(RUN, structured, "docx")
        document = Document(BytesIO(out.content))
        assert document.paragraphs[0].text == "Government Work Report"
        assert document.paragraphs[0].style.name == "Title"
        assert document.paragraphs[1].text == "I. General Requirements"
        assert document.paragraphs[1].runs[0].bold is True
        assert round(document.sections[0].top_margin.cm, 1) == 2.4
        assert not any("pipeline" in paragraph.text.lower() for paragraph in document.paragraphs)

    def test_all_formats_registered(self):
        assert set(FORMATS) == {"txt", "md", "json", "xliff", "tmx", "docx", "docx_bilingual"}
