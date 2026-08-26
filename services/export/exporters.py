"""Export Engine (§39). Formats: txt / md / json / xliff / tmx / docx /
docx_bilingual. DOCX keeps heading/paragraph/list/table structure when the
run carries parsed structure; plain runs degrade to paragraphs.

All exporters are pure functions over (run, segments) so they are unit
testable without the API layer.
"""

from __future__ import annotations

import io
import json
import re
from dataclasses import dataclass
from datetime import datetime, timezone
from xml.sax.saxutils import escape

from services.languages import language_spec, resolve_language_pair
from services.orchestrator.segmentation import infer_block_kind


@dataclass
class ExportResult:
    content: bytes
    media_type: str
    filename: str


FORMATS = ("txt", "md", "json", "xliff", "tmx", "docx", "docx_bilingual")


def _esc(text: str) -> str:
    return escape(text, {'"': "&quot;"})


def export_txt(segments: list[dict]) -> bytes:
    return "\n\n".join(s["translation"] or "" for s in segments).encode("utf-8")


def _run_languages(run: dict):
    source, target = resolve_language_pair(
        run.get("source_language"),
        run.get("target_language"),
        run.get("direction"),
    )
    return language_spec(source), language_spec(target)


def _document_title(run: dict, segments: list[dict]) -> tuple[str, int]:
    """Return a human-facing title and the number of body rows to skip."""
    if segments:
        first = segments[0]
        translated = (first.get("translation") or "").strip()
        if translated and infer_block_kind(first.get("source", ""), index=0) == "title":
            return translated[:180], 1
    summary = " ".join(str(run.get("summary") or "").split())
    _source_language, target_language = _run_languages(run)
    return (summary[:180] or f"{target_language.name_en} Translation"), 0


def _safe_stem(title: str, run_id: str) -> str:
    ascii_title = title.encode("ascii", "ignore").decode()
    stem = re.sub(r"[^A-Za-z0-9]+", "-", ascii_title).strip("-").lower()
    return stem[:64].strip("-") or f"govtrans-{run_id[:8]}"


def export_md(run: dict, segments: list[dict]) -> bytes:
    title, body_start = _document_title(run, segments)
    lines = [f"# {title}", ""]
    for index, seg in enumerate(segments[body_start:], start=body_start):
        kind = infer_block_kind(seg.get("source", ""), index=index)
        if kind == "heading":
            lines.append(f"## {seg['translation'] or ''}")
        else:
            lines.append(seg["translation"] or "")
        lines.append("")
    return "\n".join(lines).encode("utf-8")


def export_json(run: dict, segments: list[dict]) -> bytes:
    payload = {
        "run_id": run["id"],
        "direction": run["direction"],
        "source_language": _run_languages(run)[0].code,
        "target_language": _run_languages(run)[1].code,
        "status": run["status"],
        "pipeline_version": run["pipeline_version"],
        "version_pins": run["version_pins"],
        "segments": [
            {
                "idx": s["idx"],
                "source": s["source"],
                "translation": s["translation"],
                "versions": s["versions"],
            }
            for s in segments
        ],
    }
    return json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")


def export_xliff(run: dict, segments: list[dict]) -> bytes:
    source_language, target_language = _run_languages(run)
    units = []
    for seg in segments:
        units.append(
            f'    <trans-unit id="{seg["idx"]}" xml:space="preserve">\n'
            f'      <source xml:lang="{source_language.bcp47}">'
            f"{_esc(seg['source'])}</source>\n"
            f'      <target xml:lang="{target_language.bcp47}">'
            f"{_esc(seg['translation'] or '')}</target>\n"
            f"    </trans-unit>"
        )
    doc = (
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<xliff version="1.2" xmlns="urn:oasis:names:tc:xliff:document:1.2">\n'
        f'  <file original="{_esc(run["id"])}" '
        f'source-language="{source_language.bcp47}" '
        f'target-language="{target_language.bcp47}" datatype="plaintext">\n'
        "  <body>\n" + "\n".join(units) + "\n  </body>\n  </file>\n</xliff>\n"
    )
    return doc.encode("utf-8")


def export_tmx(run: dict, segments: list[dict]) -> bytes:
    source_language, target_language = _run_languages(run)
    tus = []
    created = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    for seg in segments:
        tus.append(
            f'    <tu tuid="{seg["idx"]}" creationdate="{created}">\n'
            f'      <tuv xml:lang="{source_language.bcp47}">'
            f"<seg>{_esc(seg['source'])}</seg></tuv>\n"
            f'      <tuv xml:lang="{target_language.bcp47}">'
            f"<seg>{_esc(seg['translation'] or '')}</seg></tuv>\n"
            "    </tu>"
        )
    doc = (
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<tmx version="1.4">\n'
        '  <header creationtool="GovTrans" creationtoolversion="'
        + _esc(run["pipeline_version"])
        + f'" segtype="sentence" adminlang="en-US" '
        f'srclang="{source_language.bcp47}" datatype="plaintext"/>\n'
        "  <body>\n" + "\n".join(tus) + "\n  </body>\n</tmx>\n"
    )
    return doc.encode("utf-8")


def export_docx(run: dict, segments: list[dict], *, bilingual: bool) -> bytes:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    source_language, target_language = _run_languages(run)

    fonts = {
        "zh": "Microsoft YaHei",
        "ja": "Yu Gothic",
        "ko": "Malgun Gothic",
        "hi": "Nirmala UI",
        "th": "Leelawadee UI",
    }
    source_font = fonts.get(source_language.code, "Arial")
    target_font = fonts.get(target_language.code, "Arial")

    def set_font(run_obj, name: str, size: float, *, bold: bool = False) -> None:
        run_obj.font.name = name
        run_obj.font.size = Pt(size)
        run_obj.font.bold = bold
        run_obj._element.rPr.rFonts.set(qn("w:eastAsia"), name)

    def shade(cell, fill: str) -> None:
        properties = cell._tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        properties.append(shading)

    def repeat_table_header(row) -> None:
        properties = row._tr.get_or_add_trPr()
        repeat = OxmlElement("w:tblHeader")
        repeat.set(qn("w:val"), "true")
        properties.append(repeat)

    def set_rtl(paragraph, enabled: bool) -> None:
        if not enabled:
            return
        properties = paragraph._p.get_or_add_pPr()
        bidi = OxmlElement("w:bidi")
        bidi.set(qn("w:val"), "1")
        properties.append(bidi)
        if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.4)
    section.left_margin = Cm(2.55)
    section.right_margin = Cm(2.55)

    normal = document.styles["Normal"]
    normal.font.name = target_font
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), target_font)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(8)

    title, body_start = _document_title(run, segments)
    heading = document.add_paragraph()
    heading.style = document.styles["Title"]
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_after = Pt(18)
    set_font(heading.add_run(title), target_font, 17, bold=True)
    set_rtl(heading, target_language.rtl)

    if bilingual:
        if body_start and segments:
            source_title = document.add_paragraph()
            source_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            source_title.paragraph_format.space_after = Pt(14)
            source_run = source_title.add_run(segments[0].get("source", ""))
            set_font(source_run, source_font, 10.5)
            set_rtl(source_title, source_language.rtl)
            source_run.font.color.rgb = RGBColor(90, 101, 117)

        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        hdr = table.rows[0].cells
        repeat_table_header(table.rows[0])
        labels = (
            f"{source_language.name_zh}原文",
            f"{target_language.name_zh}译文",
        )
        for cell, label, width in zip(hdr, labels, (7.4, 8.2)):
            cell.width = Cm(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            shade(cell, "25364D")
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
            header_run = paragraph.add_run(label)
            set_font(header_run, target_font, 10, bold=True)
            header_run.font.color.rgb = RGBColor(255, 255, 255)

        for seg in segments[body_start:]:
            cells = table.add_row().cells
            cells[0].width = Cm(7.4)
            cells[1].width = Cm(8.2)
            cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            source_paragraph = cells[0].paragraphs[0]
            target_paragraph = cells[1].paragraphs[0]
            source_paragraph.paragraph_format.space_after = Pt(2)
            target_paragraph.paragraph_format.space_after = Pt(2)
            set_font(source_paragraph.add_run(seg.get("source", "")), source_font, 10)
            set_font(target_paragraph.add_run(seg.get("translation") or ""), target_font, 10)
            set_rtl(source_paragraph, source_language.rtl)
            set_rtl(target_paragraph, target_language.rtl)
    else:
        for index, seg in enumerate(segments[body_start:], start=body_start):
            text = seg.get("translation") or ""
            kind = infer_block_kind(seg.get("source", ""), index=index)
            if kind == "heading":
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_before = Pt(12)
                paragraph.paragraph_format.space_after = Pt(7)
                set_font(paragraph.add_run(text), target_font, 13, bold=True)
            else:
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                if kind == "list":
                    paragraph.paragraph_format.left_indent = Cm(0.45)
                    paragraph.paragraph_format.first_line_indent = Cm(-0.2)
                set_font(paragraph.add_run(text), target_font, 11)
            set_rtl(paragraph, target_language.rtl)
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def export_run(run: dict, segments: list[dict], fmt: str) -> ExportResult:
    if fmt not in FORMATS:
        raise ValueError(f"unsupported format {fmt!r}; choose from {FORMATS}")
    title, _body_start = _document_title(run, segments)
    stem = _safe_stem(title, run["id"])
    if fmt == "txt":
        return ExportResult(export_txt(segments), "text/plain; charset=utf-8", f"{stem}.txt")
    if fmt == "md":
        return ExportResult(export_md(run, segments), "text/markdown; charset=utf-8", f"{stem}.md")
    if fmt == "json":
        return ExportResult(export_json(run, segments), "application/json", f"{stem}.json")
    if fmt == "xliff":
        return ExportResult(export_xliff(run, segments), "application/x-xliff+xml", f"{stem}.xlf")
    if fmt == "tmx":
        return ExportResult(export_tmx(run, segments), "application/x-tmx+xml", f"{stem}.tmx")
    if fmt == "docx":
        content = export_docx(run, segments, bilingual=False)
        return ExportResult(
            content,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            f"{stem}.docx",
        )
    content = export_docx(run, segments, bilingual=True)
    return ExportResult(
        content,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        f"{stem}-bilingual.docx",
    )
